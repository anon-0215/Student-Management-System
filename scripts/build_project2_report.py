from pathlib import Path
import shutil
from textwrap import wrap

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from PIL import Image, ImageDraw, ImageFont


ROOT = Path(__file__).resolve().parents[1]
OUT_DIR = ROOT / "output" / "documents"
ASSET_DIR = OUT_DIR / "project2_report_assets"
REPORT_PATH = OUT_DIR / "校园食堂信息管理系统_project_2_开发技术报告_真实截图最终版.docx"

FONT_CJK = "C:/Windows/Fonts/msyh.ttc"
FONT_CJK_BOLD = "C:/Windows/Fonts/simhei.ttf"


BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
INK = RGBColor(25, 35, 45)
MUTED = RGBColor(92, 99, 112)
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F2F4F7"


def ensure_dirs():
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    ASSET_DIR.mkdir(parents=True, exist_ok=True)


def font(size=28, bold=False):
    path = FONT_CJK_BOLD if bold else FONT_CJK
    return ImageFont.truetype(path, size=size)


def wrap_text(draw, text, font_obj, max_width):
    lines = []
    for para in text.split("\n"):
        current = ""
        for ch in para:
            candidate = current + ch
            if draw.textbbox((0, 0), candidate, font=font_obj)[2] <= max_width:
                current = candidate
            else:
                if current:
                    lines.append(current)
                current = ch
        if current:
            lines.append(current)
    return lines or [""]


def text_box(draw, xy, text, fill="#FFFFFF", outline="#2E74B5", width=3,
             radius=18, text_fill="#1F2A37", size=30, bold=False, align="center"):
    x1, y1, x2, y2 = xy
    draw.rounded_rectangle(xy, radius=radius, fill=fill, outline=outline, width=width)
    fnt = font(size, bold)
    lines = wrap_text(draw, text, fnt, x2 - x1 - 36)
    line_h = size + 10
    total_h = line_h * len(lines)
    y = y1 + ((y2 - y1) - total_h) / 2
    for line in lines:
        bbox = draw.textbbox((0, 0), line, font=fnt)
        if align == "left":
            x = x1 + 22
        else:
            x = x1 + ((x2 - x1) - (bbox[2] - bbox[0])) / 2
        draw.text((x, y), line, font=fnt, fill=text_fill)
        y += line_h


def arrow(draw, start, end, color="#2E74B5", width=5):
    draw.line([start, end], fill=color, width=width)
    ex, ey = end
    sx, sy = start
    if abs(ex - sx) >= abs(ey - sy):
        sign = 1 if ex > sx else -1
        pts = [(ex, ey), (ex - sign * 18, ey - 12), (ex - sign * 18, ey + 12)]
    else:
        sign = 1 if ey > sy else -1
        pts = [(ex, ey), (ex - 12, ey - sign * 18), (ex + 12, ey - sign * 18)]
    draw.polygon(pts, fill=color)


def save_diagrams():
    # Function module diagram
    img = Image.new("RGB", (1800, 1150), "#FFFFFF")
    d = ImageDraw.Draw(img)
    d.text((70, 45), "系统功能模块图", font=font(44, True), fill="#0B2545")
    d.text((70, 105), "project_2 采用主菜单分流，管理员使用全局函数管理全局数据，学生通过对象成员函数操作自身数据。",
           font=font(25), fill="#4B5563")
    text_box(d, (650, 160, 1150, 250), "校园学生食堂信息管理系统", "#0B2545", "#0B2545",
             text_fill="#FFFFFF", size=32, bold=True)
    text_box(d, (230, 330, 760, 420), "管理员端", "#E8EEF5", "#2E74B5", size=34, bold=True)
    text_box(d, (1040, 330, 1570, 420), "学生端", "#E8EEF5", "#2E74B5", size=34, bold=True)
    arrow(d, (760, 250), (500, 330))
    arrow(d, (1040, 250), (1300, 330))
    admin_items = [
        "学生数据管理\n查看、增加、删除、修改、查询订单",
        "菜品管理\n增删改查、搜索筛选、低库存提示",
        "订单查询\n全部订单、按学生、按菜品关键字",
        "数据统计分析\n营收、状态、时段、方式、支付、热销排行",
        "评价与投诉处理\n查看菜品评价、回复投诉建议",
    ]
    stu_items = [
        "用户入口\n注册、登录、查看/修改个人信息",
        "菜品浏览\n列表、详情、分类/关键字/价格筛选",
        "订餐支付取餐\n选菜、数量、时段、方式、支付、取餐码",
        "订单服务\n查询订单、取消订单、退款、热销排行",
        "评价反馈\n评分留言、提交投诉、查看回复",
    ]
    y = 485
    for item in admin_items:
        text_box(d, (110, y, 860, y + 95), item, "#F8FAFC", "#93C5FD", size=25, align="left")
        y += 118
    y = 485
    for item in stu_items:
        text_box(d, (940, y, 1690, y + 95), item, "#F8FAFC", "#93C5FD", size=25, align="left")
        y += 118
    out = ASSET_DIR / "function_modules.png"
    img.save(out)

    # Data structure diagram
    img = Image.new("RGB", (1800, 1120), "#FFFFFF")
    d = ImageDraw.Draw(img)
    d.text((70, 45), "对象与数据结构关系图", font=font(44, True), fill="#0B2545")
    d.text((70, 105), "学生数据使用链表组织，菜品与投诉使用 vector 保存；订单属于学生对象自身。",
           font=font(25), fill="#4B5563")
    text_box(d, (90, 210, 410, 310), "StudentList\nhead", "#0B2545", "#0B2545",
             text_fill="#FFFFFF", size=30, bold=True)
    node_positions = [(560, 195, 900, 325), (1030, 195, 1370, 325), (1500, 195, 1760, 325)]
    for i, pos in enumerate(node_positions, start=1):
        text_box(d, pos, f"Student 节点 {i}\n学号/密码/姓名/年级/专业/电话\nnext 指针", "#E8EEF5", "#2E74B5", size=23)
    arrow(d, (410, 260), (560, 260))
    arrow(d, (900, 260), (1030, 260))
    arrow(d, (1370, 260), (1500, 260))
    text_box(d, (620, 420, 1160, 550), "每个 Student 内部保存\nvector<Order> orders\n只记录该学生自己的订单", "#F8FAFC", "#94A3B8", size=26)
    arrow(d, (730, 325), (820, 420), "#94A3B8")
    text_box(d, (110, 700, 610, 860), "vector<Dish> dishes\n菜品编号、名称、分类、价格、库存、口味、营养、过敏源、评分、留言", "#F2F4F7", "#64748B", size=25)
    text_box(d, (690, 700, 1190, 860), "vector<Complaint> complaints\n投诉编号、学生、类型、内容、回复、处理状态", "#F2F4F7", "#64748B", size=25)
    text_box(d, (1270, 700, 1690, 860), "全局函数\n管理员菜单、订单查询、统计分析、热销排行", "#F2F4F7", "#64748B", size=25)
    arrow(d, (1450, 700), (1180, 550), "#64748B")
    arrow(d, (1450, 700), (500, 700), "#64748B")
    arrow(d, (1450, 700), (950, 700), "#64748B")
    d.text((110, 960), "设计重点：成员函数有 this 指针，适合维护调用对象自身数据；全局函数没有 this 指针，适合处理系统级菜单、查询与统计。",
           font=font(27, True), fill="#1F4D78")
    out = ASSET_DIR / "data_structures.png"
    img.save(out)

    # Order lifecycle flow
    img = Image.new("RGB", (1800, 1200), "#FFFFFF")
    d = ImageDraw.Draw(img)
    d.text((70, 45), "学生订餐-支付-取餐业务流程图", font=font(44, True), fill="#0B2545")
    d.text((70, 105), "核心流程由 Student 成员函数完成，订单保存在当前学生对象的 orders 向量中。",
           font=font(25), fill="#4B5563")
    steps = [
        ((130, 210, 500, 320), "学生登录\nstudentLogin 查找链表节点"),
        ((650, 210, 1030, 320), "查看菜品\nvector<Dish> 列表与详情"),
        ((1180, 210, 1600, 320), "选择菜品和数量\n校验编号与库存"),
        ((1180, 445, 1600, 555), "生成订单\n状态 Reserved\n扣减库存并生成取餐码"),
        ((650, 445, 1030, 555), "支付订单\n选择校园卡/微信/支付宝\n状态 Paid"),
        ((130, 445, 500, 555), "取餐验证\n核对订单号和取餐码\n状态 PickedUp"),
        ((130, 735, 500, 845), "评价菜品\n仅已取餐且未评价订单"),
        ((650, 735, 1030, 845), "统计/热销排行\n只统计 Paid/PickedUp"),
        ((1180, 735, 1600, 845), "异常分支\n取消订单/退款\n恢复库存"),
    ]
    for pos, label in steps:
        text_box(d, pos, label, "#F8FAFC", "#2E74B5", size=25)
    arrow(d, (500, 265), (650, 265))
    arrow(d, (1030, 265), (1180, 265))
    arrow(d, (1390, 320), (1390, 445))
    arrow(d, (1180, 500), (1030, 500))
    arrow(d, (650, 500), (500, 500))
    arrow(d, (315, 555), (315, 735))
    arrow(d, (500, 790), (650, 790))
    arrow(d, (1030, 790), (1180, 790), "#9B1C1C")
    text_box(d, (300, 990, 1500, 1100), "状态流转：Reserved（待支付） -> Paid（已支付） -> PickedUp（已取餐）；Reserved 可取消，Paid 可退款，PickedUp 不允许取消或退款。",
             "#FFF7ED", "#F59E0B", size=27)
    out = ASSET_DIR / "order_flow.png"
    img.save(out)

    # Interaction diagram
    img = Image.new("RGB", (1800, 1080), "#FFFFFF")
    d = ImageDraw.Draw(img)
    d.text((70, 45), "菜单交互与函数调用图", font=font(44, True), fill="#0B2545")
    d.text((70, 105), "系统使用多级菜单组织交互，管理员入口调用全局函数，学生入口调用 Student 成员函数。",
           font=font(25), fill="#4B5563")
    text_box(d, (690, 180, 1110, 270), "main()", "#0B2545", "#0B2545", text_fill="#FFFFFF", size=34, bold=True)
    text_box(d, (690, 340, 1110, 430), "mainMenu()", "#E8EEF5", "#2E74B5", size=32, bold=True)
    arrow(d, (900, 270), (900, 340))
    text_box(d, (180, 545, 700, 655), "管理员登录\nadminLogin()", "#F8FAFC", "#2E74B5", size=28)
    text_box(d, (1100, 545, 1620, 655), "学生登录/注册\nstudentLogin() / registerStudent()", "#F8FAFC", "#2E74B5", size=28)
    arrow(d, (800, 430), (440, 545))
    arrow(d, (1000, 430), (1360, 545))
    text_box(d, (80, 760, 800, 900), "adminMenu()\nstudentManageMenu / dishManageMenu / orderQueryMenu / dataAnalysisMenu / complaintManageMenu", "#F2F4F7", "#64748B", size=25)
    text_box(d, (1000, 760, 1720, 900), "studentMenu()\nviewDishes / placeOrder / payOrder / pickupOrder / showOrders / rateDish / submitComplaint", "#F2F4F7", "#64748B", size=25)
    arrow(d, (440, 655), (440, 760), "#64748B")
    arrow(d, (1360, 655), (1360, 760), "#64748B")
    out = ASSET_DIR / "interaction.png"
    img.save(out)


def terminal_image(name, title, lines):
    width = 1500
    padding = 48
    title_h = 82
    line_h = 36
    wrapped = []
    fnt = font(25)
    dummy = ImageDraw.Draw(Image.new("RGB", (10, 10)))
    prompt = r"PS C:\Users\Anon\Desktop\餐厅管理系统> .\project_2.exe"
    for line in [prompt, ""] + lines:
        if not line:
            wrapped.append("")
            continue
        current_lines = wrap_text(dummy, line, fnt, width - padding * 2)
        wrapped.extend(current_lines)
    height = title_h + padding + line_h * len(wrapped) + 52
    img = Image.new("RGB", (width, height), "#0C0C0C")
    d = ImageDraw.Draw(img)
    d.rectangle((0, 0, width, title_h), fill="#202020")
    d.rectangle((20, 22, 56, 58), fill="#012456", outline="#3B82F6", width=2)
    d.text((28, 24), ">_", font=font(16, True), fill="#FFFFFF")
    d.text((74, 22), f"Windows PowerShell - {title}", font=font(26, True), fill="#E5E7EB")
    control_w = 66
    labels = ["-", "□", "×"]
    for i, label in enumerate(labels):
        x1 = width - control_w * (3 - i)
        x2 = x1 + control_w
        fill = "#202020" if label != "×" else "#B91C1C"
        d.rectangle((x1, 0, x2, title_h), fill=fill)
        text_fill = "#D1D5DB" if label != "×" else "#FFFFFF"
        bbox = d.textbbox((0, 0), label, font=font(24, True))
        d.text((x1 + (control_w - (bbox[2] - bbox[0])) / 2, 24), label, font=font(24, True), fill=text_fill)
    y = title_h + 34
    for index, line in enumerate(wrapped):
        line_fill = "#9CDCFE" if index == 0 else "#F2F2F2"
        d.text((padding, y), line, font=fnt, fill=line_fill)
        y += line_h
    out = ASSET_DIR / name
    img.save(out)
    return out


def save_screenshots():
    source_dir = ASSET_DIR / "real_screenshots"
    mappings = [
        ("01_main_menu.png", "run_main_menu.png", None),
        ("02_student_menu.png", "run_student_menu.png", None),
        ("03_order_success.png", "run_order_success.png", (52, 0, 1793, 1458)),
        ("04_payment_success.png", "run_payment_success.png", None),
        ("05_pickup_success.png", "run_pickup_success.png", None),
        ("06_order_query.png", "run_order_query.png", None),
        ("07_admin_stats.png", "run_admin_stats.png", (215, 5, 1620, 1515)),
        ("08_rating.png", "run_rating.png", None),
    ]
    for src_name, dst_name, crop in mappings:
        src = source_dir / src_name
        dst = ASSET_DIR / dst_name
        if not src.exists():
            raise FileNotFoundError(f"Missing real screenshot: {src}")
        if crop is None:
            shutil.copyfile(src, dst)
        else:
            img = Image.open(src)
            img.crop(crop).save(dst)


def set_cell_text(cell, text, bold=False, color=None):
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(text)
    set_run_font(run, bold=bold, color=color)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in [("top", top), ("start", start), ("bottom", bottom), ("end", end)]:
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_table_widths(table, widths):
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    for row in table.rows:
        for idx, width in enumerate(widths):
            row.cells[idx].width = Inches(width)
            set_cell_margins(row.cells[idx])
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:type"), "dxa")
    tbl_w.set(qn("w:w"), "9360")


def set_run_font(run, name="Calibri", east="宋体", size=None, bold=None, color=None):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), east)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color is not None:
        run.font.color.rgb = color


def set_style_font(style, size, color=None, bold=None, before=None, after=None, line_spacing=None):
    style.font.name = "Calibri"
    style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    style.font.size = Pt(size)
    if color is not None:
        style.font.color.rgb = color
    if bold is not None:
        style.font.bold = bold
    if before is not None:
        style.paragraph_format.space_before = Pt(before)
    if after is not None:
        style.paragraph_format.space_after = Pt(after)
    if line_spacing is not None:
        style.paragraph_format.line_spacing = line_spacing


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("第 ")
    set_run_font(run, size=9, color=MUTED)
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_end)
    tail = paragraph.add_run(" 页")
    set_run_font(tail, size=9, color=MUTED)


def configure_document(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    set_style_font(styles["Normal"], 11, color=INK, after=6, line_spacing=1.10)
    set_style_font(styles["Heading 1"], 16, color=BLUE, bold=True, before=16, after=8)
    set_style_font(styles["Heading 2"], 13, color=BLUE, bold=True, before=12, after=6)
    set_style_font(styles["Heading 3"], 12, color=DARK_BLUE, bold=True, before=8, after=4)
    styles.add_style("CodeBlock", 1)
    set_style_font(styles["CodeBlock"], 9, color=RGBColor(30, 41, 59), after=4, line_spacing=1.0)
    styles["CodeBlock"].font.name = "Consolas"
    styles["CodeBlock"]._element.rPr.rFonts.set(qn("w:ascii"), "Consolas")
    styles["CodeBlock"]._element.rPr.rFonts.set(qn("w:hAnsi"), "Consolas")
    styles["CodeBlock"]._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")

    header_p = section.header.paragraphs[0]
    header_p.text = ""
    header_run = header_p.add_run("校园食堂信息管理系统开发技术报告 | project_2.cpp")
    set_run_font(header_run, size=9, color=MUTED)
    footer_p = section.footer.paragraphs[0]
    footer_p.text = ""
    add_page_number(footer_p)


def add_title_page(doc):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(70)
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run("校园食堂信息管理系统")
    set_run_font(r, size=26, bold=True, color=RGBColor(0, 0, 0))

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(28)
    r = p.add_run("project_2 源码开发技术报告")
    set_run_font(r, size=16, color=RGBColor(70, 70, 70))

    table = doc.add_table(rows=7, cols=2)
    set_table_widths(table, [1.75, 4.75])
    table.style = "Table Grid"
    rows = [
        ("课程名称", "2026 年春季学期编程实践课程"),
        ("项目名称", "校园学生食堂信息管理系统"),
        ("对应源码", "project_2.cpp"),
        ("开发语言", "C++ 控制台程序"),
        ("姓名", "____________"),
        ("学号", "____________"),
        ("班级", "____________"),
    ]
    for idx, (label, value) in enumerate(rows):
        set_cell_text(table.rows[idx].cells[0], label, bold=True)
        shade_cell(table.rows[idx].cells[0], LIGHT_GRAY)
        set_cell_text(table.rows[idx].cells[1], value)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(32)
    r = p.add_run("完成时间：2026 年 7 月")
    set_run_font(r, size=11, color=MUTED)
    doc.add_page_break()


def add_contents(doc):
    doc.add_heading("目录", level=1)
    entries = [
        "1. 项目概述与需求分析",
        "2. 开发环境与总体技术方案",
        "3. 系统总体设计",
        "4. 核心功能实现",
        "5. 运行结果与功能验证",
        "6. 项目特色与需求完成情况",
        "7. 不足与改进方向",
        "8. 编程心得与总结",
    ]
    for entry in entries:
        p = doc.add_paragraph(entry)
        p.paragraph_format.left_indent = Inches(0.2)
        p.paragraph_format.space_after = Pt(4)
    doc.add_page_break()


def add_table(doc, headers, rows, widths):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    set_table_widths(table, widths)
    for i, head in enumerate(headers):
        set_cell_text(table.rows[0].cells[i], head, bold=True, color=RGBColor(0, 0, 0))
        shade_cell(table.rows[0].cells[i], LIGHT_BLUE)
    for row_data in rows:
        row = table.add_row()
        for i, value in enumerate(row_data):
            set_cell_text(row.cells[i], value)
    return table


def add_code_block(doc, code):
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    set_table_widths(table, [6.5])
    cell = table.rows[0].cells[0]
    shade_cell(cell, "F8FAFC")
    cell.text = ""
    for idx, line in enumerate(code.strip("\n").splitlines()):
        p = cell.paragraphs[0] if idx == 0 else cell.add_paragraph()
        p.style = "CodeBlock"
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run(line)
        set_run_font(run, name="Consolas", east="宋体", size=8.7, color=RGBColor(30, 41, 59))


def add_figure(doc, path, caption, width=6.5):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(str(path), width=Inches(width))
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.space_after = Pt(8)
    r = cap.add_run(caption)
    set_run_font(r, size=9.5, color=MUTED)


def build_report():
    ensure_dirs()
    save_diagrams()
    save_screenshots()

    doc = Document()
    configure_document(doc)
    add_title_page(doc)
    add_contents(doc)

    doc.add_heading("1. 项目概述与需求分析", level=1)
    doc.add_paragraph(
        "本项目面向校园食堂的日常订餐与管理场景，开发一个控制台版校园学生食堂信息管理系统。"
        "系统围绕学生用餐、菜品管理、订单处理、支付结算、数据统计与反馈处理展开，目标是用较清晰的 C++ 面向对象结构体现"
        "“学生类链表”和“成员函数/全局函数分工”的课程要求。"
    )
    doc.add_paragraph(
        "项目需求文档要求至少完成用户管理、菜品管理、订餐与取餐等核心功能，并鼓励扩展支付结算、统计分析、评价反馈、投诉建议等实用功能。"
        "project_2.cpp 在 project_1 基础上继续增强，采用管理员和学生两个入口：管理员作为超级用户处理全局数据，学生登录后通过自身对象完成个人订单相关操作。"
    )
    rows = [
        ("用户管理", "学生注册、登录、查看与修改个人信息；管理员维护学生数据", "已实现"),
        ("角色权限", "管理员菜单与学生菜单分离，权限边界清晰", "已实现"),
        ("菜品管理", "菜品增删改查、详情展示、分类/关键字/价格筛选、低库存提示", "已实现"),
        ("订餐与取餐", "在线订餐、时段选择、堂食/自提/外卖、取餐码验证、防止重复取餐", "已实现"),
        ("支付结算", "校园卡、微信、支付宝支付；订单金额计算；取消/退款并恢复库存", "已实现"),
        ("统计分析", "订单状态、营收、时段、方式、支付方式和热销菜品统计", "已实现"),
        ("评价反馈", "学生评分留言、投诉建议提交，管理员查看和回复", "已实现"),
        ("文件/图形界面", "需求中的可选扩展方向，本版本重点完成控制台交互", "未纳入 project_2"),
    ]
    add_table(doc, ["需求方向", "project_2 对应设计", "完成情况"], rows, [1.45, 3.95, 1.10])

    doc.add_heading("2. 开发环境与总体技术方案", level=1)
    rows = [
        ("开发语言", "C++，使用标准库完成控制台交互和内存数据管理。"),
        ("主要头文件", "algorithm、cstdlib、iomanip、iostream、limits、string、vector。"),
        ("运行方式", "Windows 控制台程序，可通过 project_2.exe 运行，也可由 project_2.cpp 编译生成。"),
        ("核心数据结构", "StudentList 单链表保存学生对象；vector<Dish> 保存菜品；Student 内部 vector<Order> 保存个人订单；vector<Complaint> 保存投诉建议。"),
        ("总体思路", "把“属于学生个人”的订餐、支付、取餐、评价、投诉等操作写成 Student 成员函数；把管理员菜单、全局查询和统计分析写成全局函数。"),
    ]
    add_table(doc, ["项目", "说明"], rows, [1.45, 5.05])
    doc.add_paragraph(
        "该技术方案符合课程中“通过指针访问成员函数体现面向对象思想”的要求。管理员虽然能触发修改学生信息、查询订单等功能，"
        "但当操作涉及某个具体学生对象的数据时，仍通过 Student 指针调用成员函数完成，避免把所有逻辑都写成散乱的全局过程。"
    )

    doc.add_heading("3. 系统总体设计", level=1)
    doc.add_heading("3.1 功能模块设计", level=2)
    doc.add_paragraph(
        "系统采用多级菜单结构。主菜单负责分流到管理员登录、学生登录和学生注册；管理员菜单负责全局维护与统计；"
        "学生菜单负责个人用餐业务。"
    )
    add_figure(doc, ASSET_DIR / "function_modules.png", "图 1 系统功能模块图")

    doc.add_heading("3.2 对象与数据结构设计", level=2)
    doc.add_paragraph(
        "系统中最重要的数据结构是学生类链表。StudentList 保存头指针 head，并提供插入、删除、查找、遍历和统计人数等操作；"
        "Student 节点通过 next 指针相连。每个学生对象内部保存自己的 orders 向量，因此订单天然归属于具体学生。"
    )
    add_figure(doc, ASSET_DIR / "data_structures.png", "图 2 对象与数据结构关系图")
    rows = [
        ("Dish", "菜品编号、名称、分类、价格、库存、口味、营养、过敏源、评分、评分人数、评价留言", "用于菜单展示、订餐、库存扣减、评价统计"),
        ("Order", "订单编号、菜品编号、菜品名、数量、金额、时段、方式、状态、取餐码、支付方式、是否评价", "保存学生个人订餐全过程"),
        ("Complaint", "投诉编号、学生学号、学生姓名、类型、内容、回复、是否处理", "保存学生反馈与管理员回复"),
        ("Student", "学号、密码、姓名、性别、出生日期、年级、专业、电话、订单向量、next 指针", "既是用户信息节点，也是订单业务的主体对象"),
        ("StudentList", "head 指针及链表操作函数", "负责学生对象生命周期和链表管理"),
    ]
    add_table(doc, ["类型", "主要字段", "作用"], rows, [1.35, 3.05, 2.10])

    doc.add_heading("3.3 菜单交互设计", level=2)
    doc.add_paragraph(
        "为了让控制台程序便于操作，project_2 使用函数嵌套调用实现多级菜单。每级菜单均通过 readInt 限定输入范围，"
        "并在功能执行后调用 pauseScreen 暂停，方便用户阅读结果。"
    )
    add_figure(doc, ASSET_DIR / "interaction.png", "图 3 菜单交互与函数调用图")

    doc.add_heading("4. 核心功能实现", level=1)
    doc.add_heading("4.1 学生类链表与用户管理", level=2)
    doc.add_paragraph(
        "学生信息使用 StudentList 单链表保存。插入采用头插法，时间复杂度为 O(1)；查询和删除需要从 head 开始顺序遍历，"
        "时间复杂度为 O(n)。删除时使用 previous 指针维护前驱节点，若删除头节点则直接更新 head，否则修改 previous->next。"
    )
    add_code_block(
        doc,
        """
bool addStudent(Student* newStudent) {
    if (newStudent == nullptr) return false;
    if (findById(newStudent->getId()) != nullptr) return false;
    newStudent->next = head;
    head = newStudent;
    return true;
}
        """,
    )
    doc.add_paragraph(
        "管理员添加学生和学生注册都复用 createStudentFromInput 与 addStudent，避免重复实现。管理员修改学生信息时，先通过学号找到 Student 指针，"
        "再调用 student->modifyByAdmin()，体现对象成员函数对私有成员的管理。"
    )

    doc.add_heading("4.2 菜品管理与搜索筛选", level=2)
    doc.add_paragraph(
        "菜品使用 vector<Dish> 保存，管理员可新增、删除、修改菜品。新增菜品通过 getNextDishId 扫描当前最大编号并加 1，"
        "避免删除菜品后出现编号复用。搜索筛选支持关键字、分类、价格区间和只看有库存四种方式。"
    )
    add_code_block(
        doc,
        """
int getNextDishId(const vector<Dish>& dishes) {
    int maxId = 1000;
    for (const Dish& dish : dishes) {
        maxId = max(maxId, dish.id);
    }
    return maxId + 1;
}
        """,
    )
    doc.add_paragraph(
        "菜品展示分为简略信息和详细信息两层：简略信息适合列表浏览，详细信息展示口味、营养成分、过敏源、评分和评价留言。"
        "当库存低于 LOW_STOCK_LINE 时，列表会提示“库存偏低”。"
    )

    doc.add_heading("4.3 订餐、支付与取餐流程", level=2)
    doc.add_paragraph(
        "订餐流程由 Student::placeOrder 完成。函数先展示菜品并校验编号，再检查库存，随后读取订购数量、用餐时段和用餐方式。"
        "订单创建后状态为 Reserved，系统立即生成取餐码并扣减库存。取餐码由学号后四位和订单号拼接组成，便于演示和验证。"
    )
    add_figure(doc, ASSET_DIR / "order_flow.png", "图 4 学生订餐-支付-取餐业务流程图")
    add_code_block(
        doc,
        """
newOrder.orderId = nextOrderId++;
newOrder.amount = selectedDish.price * quantity;
newOrder.status = OrderStatus::Reserved;
newOrder.pickupCode = makePickupCode(studentId, newOrder.orderId);
orders.push_back(newOrder);
selectedDish.stock -= quantity;
        """,
    )
    doc.add_paragraph(
        "支付函数只允许待支付订单进入支付流程，已支付或已取餐订单不能重复支付，已取消或已退款订单不能支付。取餐函数要求订单已支付、取餐码正确，"
        "并把状态改为 PickedUp；若订单已取餐，则直接拒绝重复取餐。取消与退款函数会根据订单状态区分 Cancelled 和 Refunded，并恢复菜品库存。"
    )

    doc.add_heading("4.4 订单查询、统计分析与热销排行", level=2)
    doc.add_paragraph(
        "管理员订单查询通过遍历 StudentList，再读取每个学生对象的 getOrders() 常量引用完成。这样既能查看全部订单，也能按学生学号或菜品名称关键字筛选。"
        "统计分析函数同样遍历全体学生订单，累计订单状态、有效营收、时段、用餐方式和支付方式。"
    )
    add_code_block(
        doc,
        """
if (shouldCountRevenue(order.status)) {
    totalRevenue += order.amount;
}
if (order.mealTime.find("午餐") != string::npos) lunchCount++;
if (order.mode == "自提") takeAwayCount++;
if (order.payMethod == "校园卡") cardPayCount++;
        """,
    )
    doc.add_paragraph(
        "热销排行只统计 Paid 和 PickedUp 订单。函数使用三个并行 vector 保存菜品编号、菜品名和销量，遍历订单时先查找是否已存在该菜品，"
        "若存在则累加数量，否则新增一项，最后按销量降序排序并输出前 N 名。"
    )

    doc.add_heading("4.5 评价与投诉反馈", level=2)
    doc.add_paragraph(
        "评价功能要求订单必须已取餐且未评价，避免未消费评价和重复评价。评分更新使用增量平均值公式，不需要保存所有历史分数："
    )
    add_code_block(
        doc,
        """
dish.rating = (dish.rating * dish.ratingCount + score) / (dish.ratingCount + 1);
dish.ratingCount++;
order.rated = true;
        """,
    )
    doc.add_paragraph(
        "投诉建议使用 Complaint 结构体保存。学生提交时记录类型、内容和学生身份，默认回复为“暂无回复”、状态为未处理；管理员可以查看全部投诉并输入回复，"
        "处理后 handled 置为 true。"
    )

    doc.add_heading("5. 运行结果与功能验证", level=1)
    doc.add_paragraph(
        "以下运行结果图为在 Windows 终端中运行 project_2.exe 的真实截图，用于展示系统核心功能的实际运行效果。"
        "截图覆盖主菜单、学生登录、在线订餐、支付、取餐、订单查询、管理员统计以及菜品评价等关键流程。"
    )
    add_figure(doc, ASSET_DIR / "run_main_menu.png", "图 5 主菜单实际运行截图", width=6.2)
    add_figure(doc, ASSET_DIR / "run_student_menu.png", "图 6 学生登录与学生菜单实际运行截图", width=6.2)
    add_figure(doc, ASSET_DIR / "run_order_success.png", "图 7 在线订餐成功实际运行截图", width=6.2)
    add_figure(doc, ASSET_DIR / "run_payment_success.png", "图 8 支付成功实际运行截图", width=6.2)
    add_figure(doc, ASSET_DIR / "run_pickup_success.png", "图 9 取餐成功实际运行截图", width=6.2)
    add_figure(doc, ASSET_DIR / "run_order_query.png", "图 10 订单查询实际运行截图", width=6.2)
    add_figure(doc, ASSET_DIR / "run_admin_stats.png", "图 11 管理员统计与热销排行实际运行截图", width=6.2)
    add_figure(doc, ASSET_DIR / "run_rating.png", "图 12 菜品评价实际运行截图", width=6.2)
    rows = [
        ("主菜单与登录", "输入管理员 admin/123456 或学生 20260001/111111", "进入对应菜单", "通过"),
        ("学生订餐", "选择菜品 1003，数量 2，午餐，自提", "生成订单、扣减库存、取餐码 0001-1", "通过"),
        ("支付订单", "对订单 1 选择校园卡支付", "状态变为已支付，记录支付方式", "通过"),
        ("取餐验证", "输入订单号 1 与正确取餐码", "状态变为已取餐，不能重复领取", "通过"),
        ("评价菜品", "已取餐订单评分 5 星并留言", "更新平均评分并标记已评价", "通过"),
        ("统计分析", "管理员查看综合统计", "统计订单数、营收、时段、方式、支付和热销", "通过"),
        ("投诉处理", "学生提交投诉，管理员回复", "学生可查看处理状态与回复", "通过"),
    ]
    add_table(doc, ["测试项", "测试输入/操作", "预期与实际结果", "结论"], rows, [1.25, 2.05, 2.55, 0.65])

    doc.add_heading("6. 项目特色与需求完成情况", level=1)
    rows = [
        ("面向对象与链表结合", "学生是对象节点，StudentList 负责链表管理，学生订单属于学生对象自身。"),
        ("成员函数和全局函数分工清晰", "学生个人业务使用成员函数，管理员全局查询和统计使用全局函数，符合课程重点。"),
        ("业务闭环完整", "从订餐、支付、取餐、评价到投诉反馈，形成较完整的食堂服务流程。"),
        ("状态控制较严谨", "订单状态使用 enum class 管理，避免字符串状态混乱，并限制重复支付、重复取餐、已取餐退款等异常操作。"),
        ("统计维度较丰富", "综合统计包含营收、状态、时段、方式、支付方式和热销排行，体现数据分析扩展要求。"),
        ("交互简洁", "readInt/readDouble/readText 对输入范围和空值进行检查，多级菜单容易演示。"),
    ]
    add_table(doc, ["特色", "说明"], rows, [1.9, 4.6])
    doc.add_paragraph(
        "从需求匹配角度看，project_2 已覆盖题目 1 的核心要求，并实现题目 2 中“支付与结算、数据统计与分析、评价与反馈”等扩展功能。"
        "班级管理、文件自动存盘、食材库存流水和图形界面属于后续可继续完善的方向。"
    )

    doc.add_heading("7. 不足与改进方向", level=1)
    rows = [
        ("数据持久化", "当前 project_2 初始化数据写在代码中，程序退出后新增订单和评价不会保存。后续可加入 fstream，把学生、菜品、订单、投诉写入文件。"),
        ("角色细分", "当前只有管理员和学生两个入口，食堂员工角色被管理员功能覆盖。后续可增加员工账号，专门处理库存、出餐和取餐核销。"),
        ("库存管理", "当前库存是菜品份数库存，未细化到食材采购与消耗。后续可增加 Ingredient 和库存流水。"),
        ("统计维度", "当前统计基于本次运行内存数据，未支持按日/周/月筛选。后续可给订单增加日期字段并按时间汇总。"),
        ("界面体验", "当前为控制台菜单。若继续扩展，可使用 Qt 或 Web 前端实现图形化操作界面。"),
        ("安全性", "密码以明文字符串保存，仅适合课程演示。实际系统应使用加密哈希和权限校验。"),
    ]
    add_table(doc, ["方向", "改进说明"], rows, [1.45, 5.05])

    doc.add_heading("8. 编程心得与总结", level=1)
    doc.add_paragraph(
        "通过 project_2 的实现，我对 C++ 面向对象程序设计、链表指针操作和标准库容器的配合有了更直接的理解。"
        "学生对象不仅保存基本资料，还保存自己的订单记录，因此“谁拥有数据，谁负责操作数据”的设计思路比较清楚。"
        "管理员功能虽然从菜单发起，但在修改具体学生或查询学生订单时仍要先找到对象，再调用对象成员函数，这帮助我区分了成员函数和全局函数的适用场景。"
    )
    doc.add_paragraph(
        "在业务实现上，订单状态是系统正确性的关键。订餐后是待支付，支付后才能取餐，取餐后不能退款或重复领取；评价只能在取餐后进行。"
        "这些规则如果只靠提示文字很容易出错，因此程序用 OrderStatus 枚举和多处状态判断来保证流程闭环。"
    )
    doc.add_paragraph(
        "本系统仍然是一个课程实践级控制台程序，但已经覆盖了校园食堂管理的主要流程。后续如果加入文件持久化、班级管理和图形界面，"
        "可以继续提升它的实用性。整体来看，本项目较好地完成了题目对数据结构、面向对象思想、菜单交互和扩展功能的训练目标。"
    )

    doc.save(REPORT_PATH)
    return REPORT_PATH


if __name__ == "__main__":
    path = build_report()
    print(path)
